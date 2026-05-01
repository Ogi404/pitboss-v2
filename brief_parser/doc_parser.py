"""
DOCX brief parser.

Extracts keywords from DOCX briefs by detecting keyword sections
and parsing bullet lists or tables within them.
"""

import re
from io import BytesIO
from typing import List, Tuple, Optional
from docx import Document
from docx.table import Table

from .base import KeywordSpec


# Section header patterns that indicate keyword lists
KEYWORD_SECTION_PATTERNS = [
    r"primary\s*keywords?",
    r"main\s*keywords?",
    r"target\s*keywords?",
    r"secondary\s*keywords?",
    r"support\s*keywords?",
    r"lsi\s*keywords?",
    r"lsi\s*terms?",
    r"single[\s-]*word\s*lsi",
    r"seo\s*keywords?",
    r"keywords?\s*list",
]

# Group patterns for section headers
GROUP_PATTERNS = {
    "main": [r"primary", r"^main", r"target", r"focus"],
    "support": [r"secondary", r"support", r"related"],
    "lsi": [r"lsi", r"semantic", r"single[\s-]*word", r"long[\s-]*tail"],
}


def _detect_group_from_header(header: str) -> str:
    """Detect keyword group from section header text."""
    header = header.lower().strip()
    for group, patterns in GROUP_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, header, re.IGNORECASE):
                return group
    return "support"  # Default


def _is_keyword_section_header(text: str) -> bool:
    """Check if text is a keyword section header."""
    text = text.lower().strip()
    for pattern in KEYWORD_SECTION_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            return True
    return False


def _parse_usage_from_text(text: str) -> Tuple[Optional[int], Optional[int]]:
    """
    Try to extract usage count from keyword text.

    Examples:
        "casino bonus (3)" -> "casino bonus", (3, 3)
        "free spins 2-5" -> "free spins", (2, 5)
    """
    # Try parentheses format: "keyword (3)" or "keyword (2-5)"
    paren_match = re.search(r"\((\d+)(?:\s*[-–—]\s*(\d+))?\)\s*$", text)
    if paren_match:
        min_val = int(paren_match.group(1))
        max_val = int(paren_match.group(2)) if paren_match.group(2) else min_val
        return min_val, max_val

    # Try trailing number: "keyword 3" or "keyword 2-5"
    trailing_match = re.search(r"\s+(\d+)(?:\s*[-–—]\s*(\d+))?\s*$", text)
    if trailing_match:
        min_val = int(trailing_match.group(1))
        max_val = int(trailing_match.group(2)) if trailing_match.group(2) else min_val
        return min_val, max_val

    return None, None


def _clean_keyword(text: str) -> str:
    """Remove usage counts and clean keyword text."""
    # Remove parenthesized numbers
    text = re.sub(r"\s*\([\d\s\-–—]+\)\s*$", "", text)
    # Remove trailing numbers
    text = re.sub(r"\s+\d+(?:\s*[-–—]\s*\d+)?\s*$", "", text)
    # Remove bullet characters
    text = re.sub(r"^[\s•\-\*\u2022\u2023\u25E6\u2043\u2219]+", "", text)
    return text.strip()


def _parse_table_for_keywords(table: Table, current_group: str) -> List[KeywordSpec]:
    """Extract keywords from a DOCX table."""
    keywords = []

    if len(table.rows) < 2:
        return keywords

    # Get headers from first row
    headers = [cell.text.lower().strip() for cell in table.rows[0].cells]

    # Find keyword and usage columns
    keyword_col = None
    usage_col = None

    keyword_patterns = ["keyword", "term", "phrase", "search"]
    usage_patterns = ["usage", "count", "times", "how much", "frequency"]

    for i, header in enumerate(headers):
        if keyword_col is None:
            for pattern in keyword_patterns:
                if pattern in header:
                    keyword_col = i
                    break
        if usage_col is None:
            for pattern in usage_patterns:
                if pattern in header:
                    usage_col = i
                    break

    # Default to first column if no keyword column found
    if keyword_col is None:
        keyword_col = 0

    # Parse data rows
    for row in table.rows[1:]:
        cells = row.cells
        if keyword_col >= len(cells):
            continue

        keyword = cells[keyword_col].text.strip()
        if not keyword:
            continue

        req_min, req_max = None, None
        if usage_col is not None and usage_col < len(cells):
            usage_text = cells[usage_col].text.strip()
            req_min, req_max = _parse_usage_from_text(usage_text)
            if req_min is None:
                # Try parsing as plain number
                try:
                    num = int(usage_text)
                    req_min, req_max = num, num
                except ValueError:
                    pass

        keywords.append(KeywordSpec(
            keyword=keyword,
            group=current_group,
            required_min=req_min,
            required_max=req_max,
            source="parsed"
        ))

    return keywords


def parse_docx_brief(file_bytes: bytes) -> Tuple[List[KeywordSpec], str, List[str]]:
    """
    Parse a DOCX file and extract keyword specifications.

    Args:
        file_bytes: Raw bytes of the DOCX file

    Returns:
        Tuple of (keywords_list, raw_text, warnings)
    """
    keywords = []
    warnings = []
    raw_texts = []

    try:
        doc = Document(BytesIO(file_bytes))
    except Exception as e:
        warnings.append(f"Failed to read DOCX file: {e}")
        return [], "", warnings

    current_group = None
    in_keyword_section = False
    paragraphs_in_section = 0

    for element in doc.element.body:
        # Handle paragraphs
        if element.tag.endswith('p'):
            # Find the paragraph object
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    raw_texts.append(text)

                    if not text:
                        continue

                    # Check if this is a keyword section header
                    if _is_keyword_section_header(text):
                        in_keyword_section = True
                        current_group = _detect_group_from_header(text)
                        paragraphs_in_section = 0
                        continue

                    # Check if we've left the keyword section (new major header)
                    if in_keyword_section and paragraphs_in_section > 20:
                        # Probably left the section
                        in_keyword_section = False
                        current_group = None

                    # If in keyword section, parse as keyword
                    if in_keyword_section and current_group:
                        paragraphs_in_section += 1

                        # Skip if this looks like instructions
                        if len(text) > 100 or text.endswith(':'):
                            continue

                        keyword = _clean_keyword(text)
                        if keyword and len(keyword) < 80:
                            req_min, req_max = _parse_usage_from_text(text)

                            # For LSI without counts, default to min=1
                            if current_group == "lsi" and req_min is None:
                                req_min = 1

                            keywords.append(KeywordSpec(
                                keyword=keyword,
                                group=current_group,
                                required_min=req_min,
                                required_max=req_max,
                                source="parsed"
                            ))
                    break

        # Handle tables
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    # Store table as raw text
                    table_text = []
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        table_text.append(row_text)
                    raw_texts.append("\n".join(table_text))

                    # Parse table for keywords
                    group = current_group or "support"
                    table_keywords = _parse_table_for_keywords(table, group)
                    keywords.extend(table_keywords)
                    break

    raw_text = "\n".join(raw_texts)

    if not keywords:
        warnings.append("No keywords found in DOCX structure. Will use LLM fallback.")

    return keywords, raw_text, warnings
