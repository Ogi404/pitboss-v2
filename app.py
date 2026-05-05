"""
Pitboss v3 - iGaming Content Proofreading and SEO Analysis Tool

Flask application that orchestrates modular check modules for
article proofreading, keyword analysis, and SEO structure compliance.
"""

import os
import uuid
import concurrent.futures
from io import BytesIO
from pathlib import Path

import yaml

from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Allow HTTP for local OAuth development (remove in production)
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

from flask import Flask, render_template, request, redirect, url_for, session, flash
from flask_session import Session
from openai import OpenAI
from docx import Document

import config
from google_auth import (
    get_authorization_url,
    handle_oauth_callback,
    is_authorized,
    get_user_email,
    clear_credentials
)
from google_docs import (
    extract_doc_id,
    fetch_document,
    extract_text_with_positions,
    find_text_occurrences,
    apply_corrections as apply_google_docs_corrections,
    apply_comments as apply_google_docs_comments,
    get_document_url
)
from dataclasses import asdict
from checks import (
    CheckResult, WriterComment,
    ProofreadCheck, KeywordCheck, SEOStructureCheck,
    ComplianceCheck, FormattingCheck, ReadabilityCheck, WordCountCheck,
    ConsistencyCheck, StyleCheck, FactCheckModule
)
from brief_parser import parse_brief, parse_google_sheet
# Note: crawl module imported lazily in /run route to avoid Playwright dependency at startup


# ------------------------------------------------------------------
# Brand Config Loader (Phase 5)
# ------------------------------------------------------------------

def deep_merge(base: dict, override: dict) -> dict:
    """
    Deep merge override dict into base dict.

    - Nested dicts are merged recursively
    - Lists are replaced (not appended)
    - Scalar values are overridden

    Args:
        base: Base configuration dict
        override: Override configuration dict

    Returns:
        Merged configuration dict
    """
    result = base.copy()
    for key, value in override.items():
        if key in result and isinstance(result[key], dict) and isinstance(value, dict):
            result[key] = deep_merge(result[key], value)
        else:
            result[key] = value
    return result


def load_brand_config(brand_name: str) -> dict:
    """
    Load brand style guide configuration with inheritance from default.yaml.

    default.yaml is always loaded first as the base layer containing
    company-wide rules. Brand-specific YAML files only need to contain
    overrides - anything not specified falls through to default.

    Args:
        brand_name: Name of the brand (matches filename without .yaml)

    Returns:
        dict with merged brand configuration
    """
    brands_dir = Path(__file__).parent / "config" / "brands"

    # 1. Always load default.yaml as base layer
    default_path = brands_dir / "default.yaml"
    if default_path.exists():
        with open(default_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f) or {}
    else:
        config = {}

    # 2. If not default, merge brand-specific config on top
    if brand_name and brand_name != "default":
        brand_path = brands_dir / f"{brand_name}.yaml"
        if brand_path.exists():
            with open(brand_path, 'r', encoding='utf-8') as f:
                brand_config = yaml.safe_load(f) or {}
            config = deep_merge(config, brand_config)

    return config


def get_available_brands() -> list:
    """
    List available brand names from config/brands/*.yaml files.

    Returns:
        List of brand names (filenames without .yaml extension)
    """
    brands_dir = Path(__file__).parent / "config" / "brands"
    if not brands_dir.exists():
        return ["default"]
    brands = [f.stem for f in brands_dir.glob("*.yaml")]
    # Ensure 'default' is first if it exists
    if "default" in brands:
        brands.remove("default")
        brands.insert(0, "default")
    return brands if brands else ["default"]


app = Flask(__name__)

# Session configuration
app.secret_key = config.FLASK_SECRET_KEY
app.config['SESSION_TYPE'] = config.SESSION_TYPE
app.config['SESSION_FILE_DIR'] = config.SESSION_FILE_DIR
Session(app)


# OAuth credential check at startup
def check_oauth_config():
    print("\n=== OAuth Configuration Check ===")
    if config.GOOGLE_CLIENT_ID:
        cid = config.GOOGLE_CLIENT_ID
        print(f"GOOGLE_CLIENT_ID: {cid[:8]}...{cid[-20:]}")
    else:
        print("GOOGLE_CLIENT_ID: NOT SET")

    if config.GOOGLE_CLIENT_SECRET:
        print(f"GOOGLE_CLIENT_SECRET: {'*' * 10} (set)")
    else:
        print("GOOGLE_CLIENT_SECRET: NOT SET")

    print(f"GOOGLE_REDIRECT_URI: {config.GOOGLE_REDIRECT_URI}")
    print("=================================\n")


check_oauth_config()

# OpenAI client (uses OPENAI_API_KEY env var)
openai_client = OpenAI()


# ------------------------------------------------------------------
# File extraction helpers
# ------------------------------------------------------------------

def extract_text_from_docx(file_storage):
    """
    Takes an uploaded DOCX file and returns all the text as one big string.
    """
    file_bytes = file_storage.read()
    doc = Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


# ------------------------------------------------------------------
# Position tracking helpers for Google Docs
# ------------------------------------------------------------------

def enrich_corrections_with_positions(corrections, position_map):
    """
    Add position data to corrections for Google Docs integration.
    Finds all occurrences of original text and adds context for disambiguation.
    """
    enriched = []

    for i, corr in enumerate(corrections):
        original = corr.get('original', '')
        status = corr.get('status', 'ok')

        # Only process actual corrections
        if status != 'corrected' or not original:
            enriched.append({
                **corr,
                'id': f'corr_{i}',
                'position_found': False,
                'occurrences': []
            })
            continue

        # Find all occurrences with context
        occurrences = find_text_occurrences(original, position_map, context_chars=40)

        enriched.append({
            **corr,
            'id': f'corr_{i}',
            'position_found': len(occurrences) > 0,
            'occurrences': occurrences,
            # For single occurrence, set default position
            'start_index': occurrences[0]['start_index'] if len(occurrences) == 1 else None,
            'end_index': occurrences[0]['end_index'] if len(occurrences) == 1 else None,
        })

    return enriched


# ------------------------------------------------------------------
# Writer Comments helpers (Phase 6)
# ------------------------------------------------------------------

def collect_writer_comments(check_results: dict, position_map: dict = None) -> list:
    """
    Collect all WriterComment objects from check results.

    Args:
        check_results: Dict of check name -> CheckResult
        position_map: Optional position map for Google Docs integration

    Returns:
        List of comment dicts ready for template/session
    """
    all_comments = []

    for check_name, result in check_results.items():
        if not hasattr(result, 'comments') or not result.comments:
            continue

        for comment in result.comments:
            # Convert WriterComment to dict if needed
            if isinstance(comment, WriterComment):
                comment_dict = asdict(comment)
            else:
                comment_dict = dict(comment)

            # Enrich with position data if we have position_map and anchor_text
            if position_map and comment_dict.get('anchor_text'):
                occurrences = find_text_occurrences(
                    comment_dict['anchor_text'],
                    position_map,
                    context_chars=40
                )
                if occurrences:
                    comment_dict['anchor_start'] = occurrences[0]['start_index']
                    comment_dict['anchor_end'] = occurrences[0]['end_index']

            all_comments.append(comment_dict)

    return all_comments


# ------------------------------------------------------------------
# Scoring helpers (Phase 4)
# ------------------------------------------------------------------

def calculate_report_summary(quality_checks: dict, proofread_details: dict = None,
                              seo_keywords: list = None) -> dict:
    """
    Calculate overall score and status counts for the report summary bar.

    Returns:
        dict with:
            - overall_score: 0-100 score
            - pass_count: number of checks that passed
            - warn_count: number of checks with warnings
            - fail_count: number of checks that failed
    """
    status_counts = {'pass': 0, 'warn': 0, 'fail': 0, 'error': 0}

    # Weight each check category
    weights = {
        'proofread': 0.20,
        'keywords': 0.15,
        'seo_structure': 0.15,
        'compliance': 0.15,
        'consistency': 0.10,
        'style': 0.10,
        'readability': 0.10,
        'word_counts': 0.05,
    }

    scores = {}

    # Score from quality_checks (compliance, formatting, readability, word_counts, consistency, style)
    for name, check_data in quality_checks.items():
        if isinstance(check_data, dict):
            status = check_data.get('status', 'pass')
        else:
            # CheckResult object
            status = getattr(check_data, 'status', 'pass')

        status_counts[status] = status_counts.get(status, 0) + 1
        scores[name] = {'pass': 100, 'warn': 70, 'fail': 30, 'error': 0}.get(status, 50)

    # Score from proofread (based on error count)
    if proofread_details:
        sentences = proofread_details.get('sentences', [])
        error_count = sum(1 for s in sentences if s.get('status') == 'corrected')
        proofread_score = max(0, 100 - (error_count * 5))
        scores['proofread'] = proofread_score
        if error_count == 0:
            status_counts['pass'] += 1
        elif error_count <= 5:
            status_counts['warn'] += 1
        else:
            status_counts['fail'] += 1

    # Score from keywords (based on ok/warn/fail ratio)
    if seo_keywords:
        ok_count = sum(1 for k in seo_keywords if k.get('status', '').lower() == 'ok')
        total = len(seo_keywords)
        if total > 0:
            keyword_score = int((ok_count / total) * 100)
            scores['keywords'] = keyword_score
            if keyword_score >= 80:
                status_counts['pass'] += 1
            elif keyword_score >= 50:
                status_counts['warn'] += 1
            else:
                status_counts['fail'] += 1

    # Calculate weighted overall score
    total_weight = sum(weights.get(name, 0) for name in scores.keys())
    if total_weight > 0:
        overall_score = sum(
            scores.get(name, 50) * weights.get(name, 0)
            for name in scores.keys()
        ) / total_weight
    else:
        overall_score = 50

    return {
        'overall_score': int(overall_score),
        'pass_count': status_counts['pass'],
        'warn_count': status_counts['warn'],
        'fail_count': status_counts['fail'] + status_counts['error'],
    }


# ------------------------------------------------------------------
# Quality Checks Helper (Phase 9)
# ------------------------------------------------------------------

def _build_quality_checks_dict(compliance_result, formatting_result, readability_result,
                                word_count_result, consistency_result, style_result):
    """Build quality checks dict for templates, handling None for disabled checks."""
    quality_checks = {}
    if compliance_result:
        quality_checks['compliance'] = {
            'status': compliance_result.status,
            'summary': compliance_result.summary,
            'details': compliance_result.details,
        }
    if formatting_result:
        quality_checks['formatting'] = {
            'status': formatting_result.status,
            'summary': formatting_result.summary,
            'details': formatting_result.details,
        }
    if readability_result:
        quality_checks['readability'] = {
            'status': readability_result.status,
            'summary': readability_result.summary,
            'details': readability_result.details,
        }
    if word_count_result:
        quality_checks['word_counts'] = {
            'status': word_count_result.status,
            'summary': word_count_result.summary,
            'details': word_count_result.details,
        }
    if consistency_result:
        quality_checks['consistency'] = {
            'status': consistency_result.status,
            'summary': consistency_result.summary,
            'details': consistency_result.details,
        }
    if style_result:
        quality_checks['style'] = {
            'status': style_result.status,
            'summary': style_result.summary,
            'details': style_result.details,
        }
    return quality_checks


# ------------------------------------------------------------------
# Routes
# ------------------------------------------------------------------

@app.route("/", methods=["GET"])
def index():
    return render_template(
        "form.html",
        google_authorized=is_authorized(),
        google_email=get_user_email(),
        available_brands=get_available_brands()
    )


# ------------------------------------------------------------------
# OAuth Routes
# ------------------------------------------------------------------

@app.route("/auth/google")
def auth_google():
    """Initiate Google OAuth flow."""
    try:
        auth_url = get_authorization_url()
        print(f"[AUTH] Session keys before redirect: {list(session.keys())}")
        print(f"[AUTH] code_verifier set: {'code_verifier' in session}")
        print(f"[AUTH] Session type: {type(session)}")
        return redirect(auth_url)
    except ValueError as e:
        return f"OAuth configuration error: {e}", 500


@app.route("/auth/google/callback")
def auth_google_callback():
    """Handle Google OAuth callback."""
    print(f"[AUTH] Session keys in callback: {list(session.keys())}")
    print(f"[AUTH] code_verifier found: {'code_verifier' in session}")
    print(f"[AUTH] Session type: {type(session)}")
    try:
        handle_oauth_callback(request.url)
        return redirect(url_for('index'))
    except Exception as e:
        return f"OAuth error: {e}", 500


@app.route("/logout")
def logout():
    """Clear Google credentials and session."""
    clear_credentials()
    return redirect(url_for('index'))


@app.route("/parse-brief-tasks", methods=["POST"])
def parse_brief_tasks():
    """
    Parse brief file and return list of task names (for multi-task briefs).

    Used by AJAX to populate the task dropdown in the form.
    Only processes Excel files - DOCX briefs don't support multi-task.
    """
    from flask import jsonify
    import tempfile

    if 'brief_file' not in request.files:
        return jsonify({"tasks": [], "error": "No file uploaded"})

    file = request.files['brief_file']
    if not file.filename:
        return jsonify({"tasks": []})

    # Only Excel files support multi-task structure
    filename = file.filename.lower()
    if not (filename.endswith('.xls') or filename.endswith('.xlsx')):
        return jsonify({"tasks": []})

    try:
        file_bytes = file.read()
        file.seek(0)

        from brief_parser.sheet_parser import get_task_names_from_excel
        tasks = get_task_names_from_excel(file_bytes)
        return jsonify({"tasks": tasks})
    except Exception as e:
        return jsonify({"tasks": [], "error": str(e)})


@app.route("/run", methods=["POST"])
def run():
    """
    Main analysis pipeline.

    1. Parse inputs (article + brief)
    2. Run check modules (proofread, keywords, SEO)
    3. Enrich corrections with positions (for Google Docs)
    4. Render results
    """
    # 1) Get form fields and files
    english_variant = request.form.get("english_variant", "UK")
    source_type = request.form.get("source_type", "docx")
    article_file = request.files.get("article_file")
    brief_file = request.files.get("brief_file")
    google_doc_url = request.form.get("google_doc_url", "").strip()
    brand_name = request.form.get("brand_name", "default")

    # Brief source fields (Phase 8)
    brief_source_type = request.form.get("brief_source_type", "file")
    brief_sheet_url = request.form.get("brief_sheet_url", "").strip()

    # Fact-check fields (Phase 7)
    operator_url = request.form.get("operator_url", "").strip()
    geo_target = request.form.get("geo_target", "").strip()
    fact_categories = request.form.getlist("fact_categories")
    # Convert to integers
    fact_categories = [int(c) for c in fact_categories] if fact_categories else []

    # Multi-task brief support - get selected task name
    task_name = request.form.get("task_name", "").strip() or None

    # Check toggles (Phase 9) - which checks to run
    enabled_checks = request.form.getlist("enabled_checks")
    # If no checkboxes sent (edge case), default to all enabled
    if not enabled_checks:
        enabled_checks = [
            "proofread", "keywords", "seo_structure", "compliance",
            "formatting", "readability", "word_counts", "consistency", "style"
        ]

    # Validate brief source
    if brief_source_type == "google_sheet":
        if not is_authorized():
            return redirect(url_for('auth_google'))
        if not brief_sheet_url:
            return "No Google Sheet URL provided.", 400
    elif not brief_file:
        return "No task brief file uploaded.", 400

    # Variables for Google Docs integration
    doc_id = None
    position_map = None

    # 2) Extract article text based on source type
    if source_type == "google_doc":
        if not is_authorized():
            return redirect(url_for('auth_google'))

        if not google_doc_url:
            return "No Google Doc URL provided.", 400

        doc_id = extract_doc_id(google_doc_url)
        if not doc_id:
            return "Invalid Google Doc URL. Please provide a valid Google Docs link.", 400

        try:
            document = fetch_document(doc_id)
            # Check for auth error (token expired/revoked)
            if isinstance(document, dict) and document.get('auth_error'):
                session.pop('google_credentials', None)
                flash('Your Google session has expired. Please sign in again.', 'warning')
                return redirect(url_for('auth_google'))
            article_text_full, position_map = extract_text_with_positions(document)
        except Exception as e:
            return f"Error fetching Google Doc: {e}", 500
    else:
        # DOCX upload
        if not article_file:
            return "No article file uploaded.", 400

        try:
            article_text_full = extract_text_from_docx(article_file)
        except Exception as e:
            return f"Error reading article DOCX: {e}", 500

    # Parse brief with structured extraction (Phase 1 + Phase 8)
    try:
        if brief_source_type == "google_sheet":
            from google_sheets import extract_sheet_id
            sheet_id = extract_sheet_id(brief_sheet_url)
            if not sheet_id:
                return "Invalid Google Sheet URL. Please provide a valid Google Sheets link.", 400
            brief_data = parse_google_sheet(sheet_id, task_name=task_name)
            # Check for auth error
            if brief_data.parse_method == "auth_error":
                session.pop('google_credentials', None)
                flash('Your Google session has expired. Please sign in again.', 'warning')
                return redirect(url_for('auth_google'))
        else:
            brief_data = parse_brief(brief_file, task_name=task_name)
    except Exception as e:
        return f"Error reading task brief: {e}", 500

    # Load brand style guide configuration (Phase 5)
    brand_config = load_brand_config(brand_name)

    # 3) Initialize and run check modules
    # Phase 4: Parallel execution - Python checks sync, LLM checks parallel

    # -------------------------------------------------------------------------
    # Python-only checks (instant, run synchronously)
    # Only run checks that are enabled via toggles
    # -------------------------------------------------------------------------
    compliance_result = None
    if "compliance" in enabled_checks:
        compliance_result = ComplianceCheck().run(
            article_text=article_text_full,
            brand_config=brand_config
        )

    formatting_result = None
    if "formatting" in enabled_checks:
        formatting_result = FormattingCheck().run(
            article_text=article_text_full,
            brand_config=brand_config
        )

    readability_result = None
    if "readability" in enabled_checks:
        readability_result = ReadabilityCheck().run(
            article_text=article_text_full,
            brief_data=brief_data,
            brand_config=brand_config
        )

    word_count_result = None
    if "word_counts" in enabled_checks:
        word_count_result = WordCountCheck().run(
            article_text=article_text_full,
            brief_data=brief_data,
            brand_config=brand_config
        )

    # -------------------------------------------------------------------------
    # LLM checks (run in parallel for speed)
    # -------------------------------------------------------------------------
    def run_proofread():
        check = ProofreadCheck(client=openai_client, model=config.MODEL_FULL)
        return check.run(
            article_text=article_text_full,
            english_variant=english_variant,
            brand_config=brand_config
        )

    def run_keywords():
        check = KeywordCheck(client=openai_client, model=config.MODEL_MINI)
        return check.run(
            article_text=article_text_full,
            brief_data=brief_data,
            max_brief_chars=config.MAX_BRIEF_CHARS,
            brand_config=brand_config
        )

    def run_seo_structure():
        check = SEOStructureCheck(client=openai_client, model=config.MODEL_FULL)
        return check.run(
            article_text=article_text_full,
            brief_text=brief_data.raw_text,
            max_brief_chars=config.MAX_BRIEF_CHARS,
            max_article_chars=config.MAX_ARTICLE_CHARS,
            brand_config=brand_config
        )

    def run_consistency():
        check = ConsistencyCheck(client=openai_client, model=config.MODEL_MINI)
        return check.run(
            article_text=article_text_full,
            brand_config=brand_config
        )

    def run_style():
        check = StyleCheck(client=openai_client, model=config.MODEL_MINI)
        return check.run(
            article_text=article_text_full,
            brief_data=brief_data,
            brand_config=brand_config
        )

    # Only add enabled LLM checks
    llm_checks = {}
    if "proofread" in enabled_checks:
        llm_checks['proofread'] = run_proofread
    if "keywords" in enabled_checks:
        llm_checks['keywords'] = run_keywords
    if "seo_structure" in enabled_checks:
        llm_checks['seo_structure'] = run_seo_structure
    if "consistency" in enabled_checks:
        llm_checks['consistency'] = run_consistency
    if "style" in enabled_checks:
        llm_checks['style'] = run_style

    # -------------------------------------------------------------------------
    # Fact-Check (Phase 7) - Optional, only if operator_url provided
    # Lazy import to avoid Playwright dependency breaking app startup
    # -------------------------------------------------------------------------
    crawl_result = None
    if operator_url:
        # Pre-crawl the operator site (uses cache if available)
        geo_locale_arg = geo_target.split('(')[0].strip() if geo_target else ""
        print(f"[FACT-CHECK] Starting crawl for: {operator_url}")
        print(f"[FACT-CHECK] Geo locale: {geo_target}")
        print(f"[FACT-CHECK] Calling get_or_crawl with: root_url={operator_url}, geo_locale={geo_locale_arg}")
        try:
            from crawl import get_or_crawl
            crawl_result = get_or_crawl(
                root_url=operator_url,
                geo_locale=geo_locale_arg,
            )
        except ImportError as e:
            # Playwright not installed - skip fact-checking
            print(f"[FACT-CHECK] Unavailable (install playwright): {e}")
            crawl_result = None
        except Exception as e:
            # Other crawl error - continue without fact-check
            print(f"[FACT-CHECK] Crawl failed: {e}")
            crawl_result = None

        # Log crawl results for debugging
        if crawl_result:
            print(f"[FACT-CHECK] Crawl complete: {len(crawl_result.pages)} pages, {len(crawl_result.errors)} errors")
            print(f"[FACT-CHECK] Total tokens: {crawl_result.total_tokens}")
            if crawl_result.errors:
                print(f"[FACT-CHECK] Crawl errors: {crawl_result.errors[:5]}")  # First 5 errors
        else:
            print(f"[FACT-CHECK] Crawl returned None (no result)")

        if crawl_result and crawl_result.pages:
            def run_fact_check():
                check = FactCheckModule(client=openai_client, model=config.MODEL_FULL)
                return check.run(
                    article_text=article_text_full,
                    crawl_result=crawl_result,
                    geo_target=geo_target,
                    categories=fact_categories if fact_categories else [1, 3, 4, 6, 9, 12],
                )

            llm_checks['fact_check'] = run_fact_check

    llm_results = {}

    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        futures = {executor.submit(fn): name for name, fn in llm_checks.items()}

        for future in concurrent.futures.as_completed(futures):
            name = futures[future]
            try:
                llm_results[name] = future.result()
            except Exception as e:
                llm_results[name] = CheckResult(
                    name=name,
                    status="error",
                    summary=f"Check failed: {str(e)}",
                    details={"error": str(e)}
                )

    # Extract results from parallel execution (may be None if check was disabled)
    proofread_result = llm_results.get('proofread')
    keywords_result = llm_results.get('keywords')
    seo_result = llm_results.get('seo_structure')
    consistency_result = llm_results.get('consistency')
    style_result = llm_results.get('style')
    fact_check_result = llm_results.get('fact_check')  # May be None if not requested

    # Log fact-check result for debugging
    if operator_url:
        if fact_check_result:
            print(f"[FACT-CHECK] Result status: {fact_check_result.status}")
            print(f"[FACT-CHECK] Result summary: {fact_check_result.summary[:200] if fact_check_result.summary else 'None'}...")
        else:
            print(f"[FACT-CHECK] No fact-check result returned (check was not run or failed)")

    # -------------------------------------------------------------------------
    # Extract data for templates (handle None for disabled checks)
    # -------------------------------------------------------------------------

    # Proofreading data
    proofread_details = (proofread_result.details or {}) if proofread_result else {}
    sentences = proofread_details.get("sentences", [])
    clean_article = proofread_details.get("clean_article", "")

    # Keywords data
    keywords_details = (keywords_result.details or {}) if keywords_result else {}
    seo_keywords = keywords_details.get("keywords", [])
    keywords_parse_method = keywords_details.get("parse_method", "none")
    keywords_used_llm = keywords_details.get("used_llm_fallback", False)
    keywords_parse_warnings = keywords_details.get("parse_warnings", [])

    # SEO Structure data
    seo_details = (seo_result.details or {}) if seo_result else {}
    seo_headings = seo_details.get("headings", [])
    seo_meta = seo_details.get("meta", {})
    seo_structure_issues = seo_details.get("structure_issues", [])
    seo_content_quality = seo_details.get("content_quality", [])

    # Collect quality check results for template (only include enabled checks)
    quality_checks = {}
    if compliance_result:
        quality_checks["compliance"] = compliance_result
    if formatting_result:
        quality_checks["formatting"] = formatting_result
    if readability_result:
        quality_checks["readability"] = readability_result
    if word_count_result:
        quality_checks["word_counts"] = word_count_result
    if consistency_result:
        quality_checks["consistency"] = consistency_result
    if style_result:
        quality_checks["style"] = style_result

    # Collect writer comments from all checks (Phase 6)
    # Only include checks that were actually run
    all_check_results = {}
    if compliance_result:
        all_check_results["compliance"] = compliance_result
    if formatting_result:
        all_check_results["formatting"] = formatting_result
    if readability_result:
        all_check_results["readability"] = readability_result
    if word_count_result:
        all_check_results["word_counts"] = word_count_result
    if consistency_result:
        all_check_results["consistency"] = consistency_result
    if style_result:
        all_check_results["style"] = style_result
    if proofread_result:
        all_check_results["proofread"] = proofread_result
    if keywords_result:
        all_check_results["keywords"] = keywords_result
    if seo_result:
        all_check_results["seo_structure"] = seo_result
    # Add fact-check if available (Phase 7)
    if fact_check_result:
        all_check_results["fact_check"] = fact_check_result

    writer_comments = collect_writer_comments(all_check_results, position_map)

    # Calculate report summary (Phase 4) - only include enabled checks
    quality_checks_dict = {}
    if compliance_result:
        quality_checks_dict['compliance'] = {'status': compliance_result.status}
    if formatting_result:
        quality_checks_dict['formatting'] = {'status': formatting_result.status}
    if readability_result:
        quality_checks_dict['readability'] = {'status': readability_result.status}
    if word_count_result:
        quality_checks_dict['word_counts'] = {'status': word_count_result.status}
    if consistency_result:
        quality_checks_dict['consistency'] = {'status': consistency_result.status}
    if style_result:
        quality_checks_dict['style'] = {'status': style_result.status}

    report_summary = calculate_report_summary(
        quality_checks=quality_checks_dict,
        proofread_details=proofread_details if proofread_result else None,
        seo_keywords=seo_keywords if keywords_result else None
    )

    # 4) For Google Docs: enrich corrections with positions and store in session
    if source_type == "google_doc" and position_map:
        sentences = enrich_corrections_with_positions(sentences, position_map)

        # Generate session ID and store results
        session_id = str(uuid.uuid4())
        session[f'results_{session_id}'] = {
            'corrections': sentences,
            'doc_id': doc_id,
            'google_doc_url': google_doc_url,
            'source_type': source_type,
            'clean_article': clean_article,
            'english_variant': english_variant,
            'seo_data': {
                'headings': seo_headings,
                'keywords': seo_keywords,
                'meta': seo_meta,
                'structure_issues': seo_structure_issues,
                'content_quality': seo_content_quality
            },
            'keywords_parse_method': keywords_parse_method,
            'keywords_used_llm': keywords_used_llm,
            'keywords_parse_warnings': keywords_parse_warnings,
            'quality_checks': _build_quality_checks_dict(
                compliance_result, formatting_result, readability_result,
                word_count_result, consistency_result, style_result
            ),
            'report_summary': report_summary,
            'writer_comments': writer_comments,  # Phase 6
            'task_name': brief_data.task_name,  # Multi-task brief support
            # Phase 7: Fact-check data
            'fact_check': {
                'enabled': bool(operator_url),
                'operator_url': operator_url,
                'geo_target': geo_target,
                'categories': fact_categories,
                'result': {
                    'status': fact_check_result.status if fact_check_result else None,
                    'summary': fact_check_result.summary if fact_check_result else None,
                    'details': fact_check_result.details if fact_check_result else None,
                    'score': fact_check_result.score if fact_check_result else None,
                } if fact_check_result else None,
            },
        }

        return redirect(url_for('results', session_id=session_id))

    # 5) For DOCX: Store in session for export, then render report
    session_id = str(uuid.uuid4())
    session[f'results_{session_id}'] = {
        'corrections': sentences,
        'doc_id': None,
        'google_doc_url': None,
        'source_type': source_type,
        'clean_article': clean_article,
        'english_variant': english_variant,
        'seo_data': {
            'headings': seo_headings,
            'keywords': seo_keywords,
            'meta': seo_meta,
            'structure_issues': seo_structure_issues,
            'content_quality': seo_content_quality
        },
        'keywords_parse_method': keywords_parse_method,
        'keywords_used_llm': keywords_used_llm,
        'keywords_parse_warnings': keywords_parse_warnings,
        'quality_checks': _build_quality_checks_dict(
            compliance_result, formatting_result, readability_result,
            word_count_result, consistency_result, style_result
        ),
        'report_summary': report_summary,
        'writer_comments': writer_comments,
        'task_name': brief_data.task_name,  # Multi-task brief support
        'fact_check': {
            'enabled': bool(operator_url),
            'operator_url': operator_url,
            'geo_target': geo_target,
            'categories': fact_categories,
            'result': {
                'status': fact_check_result.status if fact_check_result else None,
                'summary': fact_check_result.summary if fact_check_result else None,
                'details': fact_check_result.details if fact_check_result else None,
                'score': fact_check_result.score if fact_check_result else None,
            } if fact_check_result else None,
        },
    }

    return render_template(
        "proofread_report.html",
        session_id=session_id,
        sentences=sentences,
        clean_article=clean_article,
        english_variant=english_variant,
        seo_headings=seo_headings,
        seo_keywords=seo_keywords,
        seo_meta=seo_meta,
        seo_structure_issues=seo_structure_issues,
        seo_content_quality=seo_content_quality,
        keywords_parse_method=keywords_parse_method,
        keywords_used_llm=keywords_used_llm,
        keywords_parse_warnings=keywords_parse_warnings,
        quality_checks=_build_quality_checks_dict(
            compliance_result, formatting_result, readability_result,
            word_count_result, consistency_result, style_result
        ),
        report_summary=report_summary,
        writer_comments=writer_comments,  # Phase 6
        source_type=source_type,
        task_name=brief_data.task_name,  # Multi-task brief support
        # Phase 7: Fact-check data
        fact_check={
            'enabled': bool(operator_url),
            'operator_url': operator_url,
            'geo_target': geo_target,
            'categories': fact_categories,
            'result': {
                'status': fact_check_result.status if fact_check_result else None,
                'summary': fact_check_result.summary if fact_check_result else None,
                'details': fact_check_result.details if fact_check_result else None,
                'score': fact_check_result.score if fact_check_result else None,
            } if fact_check_result else None,
        },
    )


# ------------------------------------------------------------------
# Results and Apply Routes (Google Docs)
# ------------------------------------------------------------------

@app.route("/results/<session_id>")
def results(session_id):
    """Display interactive results with checkboxes for Google Docs corrections."""
    result_data = session.get(f'results_{session_id}')
    if not result_data:
        return "Results not found or session expired. Please run the analysis again.", 404

    return render_template(
        "results.html",
        session_id=session_id,
        corrections=result_data['corrections'],
        doc_id=result_data['doc_id'],
        google_doc_url=result_data['google_doc_url'],
        source_type=result_data['source_type'],
        clean_article=result_data['clean_article'],
        english_variant=result_data['english_variant'],
        seo_headings=result_data['seo_data']['headings'],
        seo_keywords=result_data['seo_data']['keywords'],
        seo_meta=result_data['seo_data']['meta'],
        seo_structure_issues=result_data['seo_data']['structure_issues'],
        seo_content_quality=result_data['seo_data']['content_quality'],
        keywords_parse_method=result_data.get('keywords_parse_method', 'none'),
        keywords_used_llm=result_data.get('keywords_used_llm', False),
        keywords_parse_warnings=result_data.get('keywords_parse_warnings', []),
        quality_checks=result_data.get('quality_checks', {}),
        report_summary=result_data.get('report_summary', {}),
        writer_comments=result_data.get('writer_comments', []),  # Phase 6
        fact_check=result_data.get('fact_check', {}),  # Phase 7
        task_name=result_data.get('task_name'),  # Multi-task brief support
    )


@app.route("/apply-corrections", methods=["POST"])
def apply_corrections_route():
    """Apply selected corrections to the Google Doc."""
    session_id = request.form.get("session_id")
    selected_ids = request.form.getlist("corrections[]")

    if not session_id:
        return "Missing session ID.", 400

    result_data = session.get(f'results_{session_id}')
    if not result_data:
        return "Session expired. Please run the analysis again.", 404

    if result_data['source_type'] != 'google_doc':
        return "Can only apply corrections to Google Docs.", 400

    if not is_authorized():
        return redirect(url_for('auth_google'))

    doc_id = result_data['doc_id']
    all_corrections = result_data['corrections']

    # Handle occurrence selection for duplicates
    # Form data comes as: corrections[] = "corr_5" or "corr_5_occ_2"
    to_apply = []
    skipped = []

    for corr in all_corrections:
        corr_id = corr.get('id')
        if corr.get('status') != 'corrected':
            continue

        # Check if this correction was selected
        # Could be "corr_5" (use default/first) or "corr_5_occ_1" (specific occurrence)
        selected_occurrence = None

        for sel_id in selected_ids:
            if sel_id == corr_id:
                # Simple selection - use first/only occurrence
                selected_occurrence = 0
                break
            elif sel_id.startswith(f"{corr_id}_occ_"):
                # Specific occurrence selected
                try:
                    occ_num = int(sel_id.split("_occ_")[1])
                    selected_occurrence = occ_num - 1  # Convert to 0-based
                    break
                except (ValueError, IndexError):
                    pass

        if selected_occurrence is None:
            continue

        occurrences = corr.get('occurrences', [])
        if not occurrences:
            skipped.append({
                'original': corr.get('original', ''),
                'reason': 'Text not found in document'
            })
            continue

        if selected_occurrence >= len(occurrences):
            selected_occurrence = 0

        occ = occurrences[selected_occurrence]
        to_apply.append({
            'start_index': occ['start_index'],
            'end_index': occ['end_index'],
            'corrected': corr.get('corrected', ''),
            'original': corr.get('original', '')
        })

    if not to_apply and not skipped:
        return "No corrections selected.", 400

    # Apply corrections to Google Doc
    applied_count = 0
    failed = []

    if to_apply:
        try:
            result = apply_google_docs_corrections(doc_id, to_apply)
            # Check for auth error (token expired/revoked)
            if isinstance(result, dict) and result.get('auth_error'):
                session.pop('google_credentials', None)
                flash('Your Google session has expired. Please sign in again.', 'warning')
                return redirect(url_for('auth_google'))
            applied_count = result.get('applied_count', 0)
            failed = result.get('failed', [])
        except Exception as e:
            failed.append({'error': str(e)})

    return render_template(
        "apply_confirm.html",
        applied_count=applied_count,
        skipped=skipped,
        failed=failed,
        google_doc_url=get_document_url(doc_id)
    )


@app.route("/apply-comments", methods=["POST"])
def apply_comments_route():
    """
    Apply selected writer comments to the Google Doc.

    Phase 6: Writer Comments System - applies editorial comments
    that require writer attention.
    """
    session_id = request.form.get("session_id")
    selected_ids = request.form.getlist("comment_ids")

    if not session_id:
        return "Missing session ID.", 400

    result_data = session.get(f'results_{session_id}')
    if not result_data:
        return "Session expired. Please run the analysis again.", 404

    if result_data['source_type'] != 'google_doc':
        return "Can only apply comments to Google Docs.", 400

    if not is_authorized():
        return redirect(url_for('auth_google'))

    doc_id = result_data['doc_id']
    all_comments = result_data.get('writer_comments', [])

    if not selected_ids:
        return "No comments selected.", 400

    # Build list of comments to apply, including any user edits
    comments_to_apply = []
    for comment in all_comments:
        if comment.get('id') in selected_ids:
            # Get the (possibly edited) comment text from the form
            edited_text = request.form.get(
                f"comment_text_{comment['id']}",
                comment.get('comment_text', '')
            )
            comments_to_apply.append({
                'content': edited_text,
                'anchor_start': comment.get('anchor_start'),
                'anchor_end': comment.get('anchor_end'),
            })

    if not comments_to_apply:
        return "No valid comments to apply.", 400

    # Apply comments to Google Doc
    try:
        result = apply_google_docs_comments(doc_id, comments_to_apply)

        # Check for auth error
        if isinstance(result, dict) and result.get('auth_error'):
            session.pop('google_credentials', None)
            flash('Your Google session has expired. Please sign in again.', 'warning')
            return redirect(url_for('auth_google'))

        applied_count = result.get('applied_count', 0)
        failed = result.get('failed', [])

    except Exception as e:
        applied_count = 0
        failed = [{'error': str(e)}]

    return render_template(
        "apply_confirm.html",
        applied_count=applied_count,
        skipped=[],
        failed=failed,
        google_doc_url=get_document_url(doc_id),
        comment_mode=True  # Flag to adjust messaging in template
    )


# ------------------------------------------------------------------
# Report Export Route (Phase 9)
# ------------------------------------------------------------------

@app.route("/export/<session_id>")
def export_report(session_id):
    """Generate and download report as DOCX."""
    from flask import send_file

    result_data = session.get(f'results_{session_id}')
    if not result_data:
        return "Results not found or session expired.", 404

    # Build DOCX document
    doc = Document()
    doc.add_heading('Proofreading & SEO Report', 0)

    # Report summary
    report_summary = result_data.get('report_summary', {})
    if report_summary:
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(f"Overall Score: {report_summary.get('overall_score', 'N/A')}%")
        doc.add_paragraph(
            f"Pass: {report_summary.get('pass_count', 0)} | "
            f"Warnings: {report_summary.get('warn_count', 0)} | "
            f"Fails: {report_summary.get('fail_count', 0)}"
        )

    # Proofreading corrections
    corrections = result_data.get('corrections', [])
    if corrections:
        actual_corrections = [c for c in corrections if c.get('status') == 'corrected']
        if actual_corrections:
            doc.add_heading('Proofreading Corrections', level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '#'
            hdr_cells[1].text = 'Original'
            hdr_cells[2].text = 'Corrected'

            for i, corr in enumerate(actual_corrections, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = corr.get('original', '')
                row_cells[2].text = corr.get('corrected', '')

    # Keywords analysis
    seo_data = result_data.get('seo_data', {})
    keywords = seo_data.get('keywords', [])
    if keywords:
        doc.add_heading('Keyword Analysis', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Keyword'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Required'
        hdr_cells[3].text = 'Status'

        for kw in keywords:
            row_cells = table.add_row().cells
            row_cells[0].text = kw.get('keyword', '')
            row_cells[1].text = str(kw.get('count', 0))
            req = kw.get('required', '')
            row_cells[2].text = str(req) if req else '-'
            row_cells[3].text = kw.get('status', '')

    # Quality checks
    quality_checks = result_data.get('quality_checks', {})
    if quality_checks:
        doc.add_heading('Quality Checks', level=1)
        for check_name, check_data in quality_checks.items():
            status = check_data.get('status', 'unknown')
            summary = check_data.get('summary', '')
            doc.add_paragraph(f"{check_name.replace('_', ' ').title()}: {status.upper()}")
            if summary:
                doc.add_paragraph(f"  {summary}", style='List Bullet')

    # Writer comments
    writer_comments = result_data.get('writer_comments', [])
    if writer_comments:
        doc.add_heading('Writer Comments', level=1)
        for comment in writer_comments:
            severity = comment.get('severity', 'info')
            text = comment.get('comment_text', '')
            anchor = comment.get('anchor_text', '')
            doc.add_paragraph(f"[{severity.upper()}] {text}")
            if anchor:
                doc.add_paragraph(f"  Context: \"{anchor[:50]}...\"", style='List Bullet')

    # Fact-check results
    fact_check = result_data.get('fact_check', {})
    if fact_check.get('enabled') and fact_check.get('result'):
        doc.add_heading('Fact-Check Results', level=1)
        fc_result = fact_check['result']
        doc.add_paragraph(f"Status: {fc_result.get('status', 'unknown').upper()}")
        if fc_result.get('summary'):
            doc.add_paragraph(fc_result['summary'])

    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'report_{session_id[:8]}.docx'
    )


if __name__ == "__main__":
    app.run(debug=True, port=8001, use_reloader=False)
